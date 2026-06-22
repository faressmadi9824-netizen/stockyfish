"""
Shock & Deviation Dashboard
---------------------------
Cap-tier risk model with optional sector orthogonalization:

    r_stock = α + β_mkt · r_capindex + (β_sector · sector_resid) + β_vix · Δlog(VIX) + ε

Cap-tier benchmark is auto-routed by the stock's market cap:
    < $2B    → IWM (Russell 2000)
    $2B–10B  → IJH (S&P MidCap 400)
    > $10B   → SPY (S&P 500)

Optional sector factor is the sector ETF return residualized against the
cap-tier benchmark — so β_mkt and β_sector stay independently interpretable.

Reports:
  - Annualized downside / upside semi-deviation
  - Expected Shortfall at 5% (daily) — kept honest, not annualized
  - Max drawdown
  - Partial market β · partial sector β (optional) · partial VIX β
  - Idiosyncratic volatility (annualized residual σ)
  - Newey-West HAC standard errors and t-stats throughout
  - Rolling 60-day partial betas
  - Drawdown curve, return distribution, worst/best 5 days

Run:
    pip install -r requirements.txt
    streamlit run shock_dashboard.py
"""

from __future__ import annotations

import io
import math
import os
import re
import time
from concurrent.futures import ThreadPoolExecutor
from datetime import date, timedelta

import numpy as np
import pandas as pd
import requests
import statsmodels.api as sm
from statsmodels.regression.rolling import RollingOLS
import streamlit as st
import streamlit.components.v1 as components
import yfinance as yf

# yfinance caches timezone data to disk. Its default location (~/.cache) isn't
# writable on some hosts (e.g. Streamlit Community Cloud), which can make the
# .info / quoteSummary endpoint fail even though price downloads still work —
# the usual cause of "chart loads but all fundamentals are blank" in the cloud.
# Point the cache at /tmp so fundamentals load there too.
try:
    os.makedirs("/tmp/py-yfinance", exist_ok=True)
    yf.set_tz_cache_location("/tmp/py-yfinance")
except Exception:
    pass

from design_renderer import (build_html, to_js_data, to_js_fundamentals,
                             to_js_industry, to_js_tape)
from macro_data import to_js_macro
from quant_signals import (build_quant_payload, extract_earnings_drivers,
                           fetch_credit_metrics, fetch_fundamental_history,
                           fetch_sec_filing_sections)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
TRADING_DAYS = 252

# (ETF, label, lower_cap, upper_cap)
CAP_TIERS = {
    "Large": ("SPY", "S&P 500",        10e9, float("inf")),
    "Mid":   ("IJH", "S&P MidCap 400", 2e9,  10e9),
    "Small": ("IWM", "Russell 2000",   0,    2e9),
}

# yfinance sector → SPDR sector ETF
SECTOR_ETF = {
    "Technology":             "XLK",
    "Financial Services":     "XLF",
    "Healthcare":             "XLV",
    "Consumer Cyclical":      "XLY",
    "Consumer Defensive":     "XLP",
    "Industrials":            "XLI",
    "Energy":                 "XLE",
    "Utilities":              "XLU",
    "Real Estate":            "XLRE",
    "Communication Services": "XLC",
    "Basic Materials":        "XLB",
}
SECTOR_ETF_NAMES = {v: k for k, v in SECTOR_ETF.items()}
ALL_SECTOR_ETFS = list(SECTOR_ETF.values())
ALL_BENCHMARK_ETFS = [v[0] for v in CAP_TIERS.values()]

# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="Shock & Deviation Dashboard",
    page_icon="📊",
    layout="wide",
)

# Hide all Streamlit chrome — design iframe fills the page
_CHROME_CSS = (
    "[data-testid='stHeader']{display:none!important}"
    "[data-testid='stToolbar']{display:none!important}"
    "[data-testid='stDecoration']{display:none!important}"
    "#MainMenu{display:none!important}"
    "footer{display:none!important}"
    ".block-container{padding:0!important;max-width:100%!important;margin:0!important}"
    ".stApp{background:#f5f1e8!important}"
    "section[data-testid='stSidebar']{background:#faf7ef!important;border-right:1px solid #d8d1bd!important}"
)
try:
    st.html(f"<style>{_CHROME_CSS}</style>")
except AttributeError:
    st.markdown(f"<style>{_CHROME_CSS}</style>", unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Inputs
# ---------------------------------------------------------------------------
with st.sidebar:
    st.header("Inputs")
    _default_ticker = st.query_params.get("ticker", "AAPL")
    ticker = st.text_input("Ticker", value=_default_ticker).strip().upper()
    if ticker:
        st.query_params["ticker"] = ticker
    years = st.selectbox("Lookback (years)", options=[1, 3, 5, 10, 15, 20], index=2)

    st.markdown("**Cap-tier benchmark**")
    benchmark_choice = st.radio(
        "benchmark",
        options=["Auto", "Large (SPY)", "Mid (IJH)", "Small (IWM)"],
        index=0,
        label_visibility="collapsed",
        help=(
            "SPY is megacap-dominated (top 10 names ≈ 35% of weight). "
            "Auto-routes to IWM for caps < $2B, IJH for $2B–$10B, SPY for $10B+. "
            "Using SPY for a small cap overstates true market beta."
        ),
    )

    st.markdown("**Sector factor**")
    sector_mode = st.radio(
        "sector mode",
        options=["Off", "Auto-detect", "Manual"],
        index=0,
        label_visibility="collapsed",
        help=(
            "Adds the sector ETF residualized against the benchmark (Frisch-Waugh). "
            "β_mkt stays unchanged; β_sector captures the incremental sector effect. "
            "Most useful for semis, biotech, energy, REITs."
        ),
    )
    manual_sector = None
    if sector_mode == "Manual":
        manual_sector = st.selectbox(
            "Pick sector ETF",
            options=ALL_SECTOR_ETFS,
            format_func=lambda etf: f"{etf} — {SECTOR_ETF_NAMES[etf]}",
        )

    st.markdown("**Semi-deviation target**")
    target_choice = st.radio(
        "semi-dev target",
        options=["Daily risk-free (^IRX)", "Zero", "Custom"],
        index=0,
        label_visibility="collapsed",
        help=(
            "Returns below this threshold count as downside. "
            "^IRX (3-month T-bill / 252) is theoretically correct; zero is common in practice. "
            "Squared semi-deviations sum to total variance around target."
        ),
    )
    custom_target = 0.0
    if target_choice == "Custom":
        custom_target = st.number_input(
            "Custom daily target return", value=0.0001, step=0.0001, format="%.5f"
        )

    st.markdown("---")
    st.markdown("**Qualitative Research**")
    st.caption(
        "Auto-extracted from SEC EDGAR filings (10-K Items 1 / 1A / 7 + "
        "latest earnings 8-K). No API tokens required."
    )

    st.markdown("---")
    st.markdown("**Industry Analysis**")
    industry_peers_raw = st.text_input(
        "Peer tickers",
        value="",
        placeholder="blank = auto by industry",
        help=(
            "Comma-separated competitors to compare against on the Industry "
            "Analysis tab. Leave blank to auto-pick true competitors by the "
            "company's industry (e.g. ASML vs AMAT, LRCX, KLAC — not big tech)."
        ),
    )

    st.markdown("---")
    st.caption(
        "Data: Yahoo Finance · Newey-West HAC SEs, lags = 5 · 252 trading days/year."
    )

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def detect_cap_tier(market_cap) -> str:
    if market_cap is None or (isinstance(market_cap, float) and math.isnan(market_cap)):
        return "Large"
    for tier, (_, _, lo, hi) in CAP_TIERS.items():
        if lo <= market_cap < hi:
            return tier
    return "Large"


@st.cache_data(show_spinner=False, ttl=60 * 60 * 4)
def _latest_book_equity(ticker: str):
    """Most recent total stockholders' equity from the balance sheet, in the
    company's financial-reporting currency. Returns None if unavailable."""
    rows = ("Stockholders Equity", "Total Stockholder Equity",
            "Common Stock Equity", "Stockholders Equity Including Minority Interest")
    try:
        t = yf.Ticker(ticker)
        for bs in (t.quarterly_balance_sheet, t.balance_sheet):
            if bs is None or getattr(bs, "empty", True):
                continue
            for row in rows:
                if row in bs.index:
                    s = bs.loc[row].dropna()
                    if not s.empty:
                        v = float(s.iloc[0])
                        if v != 0:
                            return v
    except Exception:
        pass
    return None


@st.cache_data(show_spinner=False, ttl=60 * 60 * 4)
def _fx_rate(base: str | None, quote: str | None):
    """Units of `quote` per 1 unit of `base` (e.g. EUR→USD ≈ 1.08)."""
    if not base or not quote or base == quote:
        return 1.0
    try:
        h = yf.Ticker(f"{base}{quote}=X").history(period="5d")
        if not h.empty:
            close = h["Close"].dropna()
            if not close.empty:
                return float(close.iloc[-1])
    except Exception:
        pass
    return None


def _price_to_book(market_cap, equity_fin, fx) -> float | None:
    """P/B = market cap / book equity, with equity converted to the trading
    currency. Pure (no I/O) so it's unit-testable."""
    if not market_cap or market_cap <= 0 or equity_fin is None or fx is None:
        return None
    equity_trade = equity_fin * fx
    if equity_trade <= 0:
        return None
    return market_cap / equity_trade


def recompute_price_to_book(ticker: str, info: dict, market_cap) -> float | None:
    """Currency-consistent P/B from balance-sheet equity. None if not derivable."""
    equity = _latest_book_equity(ticker)
    if equity is None:
        return None
    fin = info.get("financialCurrency") or info.get("currency")
    trade = info.get("currency") or "USD"
    fx = _fx_rate(fin, trade)
    return _price_to_book(market_cap, equity, fx)


def _fetch_info(ticker: str) -> dict:
    """Single yfinance .info fetch.

    IMPORTANT: do NOT retry here. get_ticker_info runs for the focus ticker
    plus every industry peer (concurrently), so retrying .info multiplies the
    number of Yahoo quoteSummary requests per page load and *causes* the very
    per-IP rate-limiting that blanks the fundamentals on Streamlit Cloud. One
    call per ticker, cached 4h by get_ticker_info, keeps request volume low.
    """
    try:
        return yf.Ticker(ticker).info or {}
    except Exception:
        return {}


@st.cache_data(show_spinner=False, ttl=60 * 60 * 4)
def get_ticker_info(ticker: str) -> dict:
    try:
        info = _fetch_info(ticker)
        if not info:
            return {}
        market_cap = info.get("marketCap")
        raw_ev     = info.get("enterpriseValue")
        # Sanity-check EV: yfinance occasionally returns garbage values for
        # non-US stocks (e.g. ASML showing 37T instead of ~$670B).
        # If EV is more than 20× market cap, fall back to the standard
        # approximation: EV = marketCap + totalDebt − totalCash.
        ev_was_bad = False
        if raw_ev and market_cap and market_cap > 0 and raw_ev > market_cap * 20:
            total_debt = info.get("totalDebt") or 0
            total_cash = info.get("totalCash") or 0
            raw_ev = market_cap + total_debt - total_cash
            ev_was_bad = True

        # Recalculate EV/EBITDA and EV/Revenue from corrected EV when needed.
        raw_ev_ebitda  = info.get("enterpriseToEbitda")
        raw_ev_revenue = info.get("enterpriseToRevenue")
        if ev_was_bad and raw_ev:
            ebitda   = info.get("ebitda")
            revenue  = info.get("totalRevenue")
            raw_ev_ebitda  = (raw_ev / ebitda)  if ebitda  and ebitda  > 0 else None
            raw_ev_revenue = (raw_ev / revenue) if revenue and revenue > 0 else None

        # Sanity-check Price/Book: yfinance mis-scales priceToBook for some ADRs
        # (e.g. ASML shows ~1600× instead of ~20×) due to a currency/units
        # mismatch between the USD price and the reported book value. When the
        # value is missing or implausibly large, recompute P/B = market cap /
        # book equity, converting equity into the trading currency.
        price_to_book = info.get("priceToBook")
        if price_to_book is None or price_to_book <= 0 or price_to_book > 100:
            pb_calc = recompute_price_to_book(ticker, info, market_cap)
            if pb_calc is not None and pb_calc > 0:
                price_to_book = pb_calc
            elif price_to_book is not None and price_to_book > 100:
                # Bad value we couldn't repair — hide it rather than show garbage.
                price_to_book = None
        if price_to_book is not None and price_to_book <= 0:
            price_to_book = None  # negative book equity → P/B not meaningful

        return {
            # Identity
            "longBusinessSummary": info.get("longBusinessSummary"),
            "marketCap":       market_cap,
            "enterpriseValue": raw_ev,
            "sector":          info.get("sector"),
            "industry":        info.get("industry"),
            "longName":        info.get("longName") or info.get("shortName"),
            "exchange":        info.get("fullExchangeName") or info.get("exchange"),
            "currency":        info.get("currency", "USD"),
            "currentPrice":    info.get("currentPrice") or info.get("regularMarketPrice"),
            # Valuation
            "trailingPE":   info.get("trailingPE"),
            "forwardPE":    info.get("forwardPE"),
            "pegRatio":     info.get("pegRatio"),
            "priceToBook":  price_to_book,
            "priceToSales": info.get("priceToSalesTrailing12Months"),
            "evToEbitda":   raw_ev_ebitda,
            "evToRevenue":  raw_ev_revenue,
            # EPS
            "trailingEps":   info.get("epsTrailingTwelveMonths"),
            "forwardEps":    info.get("epsForward") or info.get("forwardEps"),
            "epsCurrentYear": info.get("epsCurrentYear"),
            # Growth
            "revenueGrowth":          info.get("revenueGrowth"),
            "earningsGrowth":         info.get("earningsGrowth"),
            "earningsQuarterlyGrowth": info.get("earningsQuarterlyGrowth"),
            # Profitability
            "grossMargins":     info.get("grossMargins"),
            "operatingMargins": info.get("operatingMargins"),
            "profitMargins":    info.get("profitMargins"),
            "ebitdaMargins":    info.get("ebitdaMargins"),
            "returnOnEquity":   info.get("returnOnEquity"),
            "returnOnAssets":   info.get("returnOnAssets"),
            # Balance sheet
            "debtToEquity": info.get("debtToEquity"),
            "currentRatio": info.get("currentRatio"),
            "quickRatio":   info.get("quickRatio"),
            # Dividend
            "dividendYield": info.get("dividendYield"),
            "payoutRatio":   info.get("payoutRatio"),
            # Technical / market
            "fiftyTwoWeekHigh":    info.get("fiftyTwoWeekHigh"),
            "fiftyTwoWeekLow":     info.get("fiftyTwoWeekLow"),
            "fiftyDayAverage":     info.get("fiftyDayAverage"),
            "twoHundredDayAverage": info.get("twoHundredDayAverage"),
            "beta":                info.get("beta"),
            "weekChange52":        info.get("52WeekChange"),
            # Analyst consensus (Refinitiv)
            "targetMeanPrice":          info.get("targetMeanPrice"),
            "targetHighPrice":          info.get("targetHighPrice"),
            "targetLowPrice":           info.get("targetLowPrice"),
            "numberOfAnalystOpinions":  info.get("numberOfAnalystOpinions"),
            "recommendationKey":        info.get("recommendationKey"),
            "averageAnalystRating":     info.get("averageAnalystRating"),
        }
    except Exception:
        return {}


_MARKET_SYMS = ALL_BENCHMARK_ETFS + ALL_SECTOR_ETFS + ["^VIX", "^VIX9D", "^VIX3M", "^IRX", "QQQ"]


@st.cache_data(show_spinner=False, ttl=60 * 60 * 4)
def load_market_prices(years: int, syms: tuple = ()) -> pd.DataFrame:
    """ETFs, VIX (+9D/3M), IRX — same for every ticker, cached 4 hours.

    `syms` is part of the cache key so adding symbols busts stale caches."""
    syms = list(syms) or _MARKET_SYMS
    end = date.today()
    start = end - timedelta(days=int(years * 365.25) + 10)
    raw = yf.download(
        syms,
        start=start.isoformat(),
        end=end.isoformat(),
        auto_adjust=True,
        progress=False,
        group_by="ticker",
        threads=True,
    )
    closes: dict[str, pd.Series] = {}
    for sym in syms:
        try:
            closes[sym] = raw[sym]["Close"]
        except (KeyError, TypeError):
            try:
                closes[sym] = raw["Close"][sym]
            except Exception:
                continue
    df = pd.DataFrame(closes).dropna(how="all")
    df.index = pd.to_datetime(df.index)
    return df


@st.cache_data(show_spinner=False, ttl=60 * 30)
def load_ticker_price(ticker: str, years: int) -> pd.Series:
    """Single ticker price history — re-fetched per ticker, cached 30 min."""
    end = date.today()
    start = end - timedelta(days=int(years * 365.25) + 10)
    raw = yf.download(
        ticker,
        start=start.isoformat(),
        end=end.isoformat(),
        auto_adjust=True,
        progress=False,
    )
    try:
        s = raw["Close"]
        if isinstance(s, pd.DataFrame):
            s = s.iloc[:, 0]
    except Exception:
        return pd.Series(dtype=float, name=ticker)
    s.index = pd.to_datetime(s.index)
    s.name = ticker
    return s


@st.cache_data(show_spinner=False, ttl=60 * 30)
def load_prices(ticker: str, years: int) -> pd.DataFrame:
    market = load_market_prices(years, tuple(_MARKET_SYMS))
    ticker_px = load_ticker_price(ticker, years)
    df = market.copy()
    df[ticker] = ticker_px.reindex(df.index)
    return df.dropna(how="all")


@st.cache_data(show_spinner=False, ttl=60)
def get_live_quote(ticker: str) -> dict:
    """Intraday last price + previous close via fast_info. Refreshes every ~1 min."""
    try:
        fi = yf.Ticker(ticker).fast_info
        return {
            "lastPrice":     float(fi.last_price)     if fi.last_price     else None,
            "previousClose": float(fi.previous_close) if fi.previous_close else None,
        }
    except Exception:
        return {}


@st.cache_data(show_spinner=False, ttl=60)
def get_live_tape_quotes(syms: tuple) -> dict:
    """Live last price + previous close for multiple tickers. Refreshes every ~1 min."""
    def _fetch(sym):
        try:
            fi = yf.Ticker(sym).fast_info
            return sym, {
                "last": float(fi.last_price     or 0),
                "prev": float(fi.previous_close or fi.last_price or 0),
            }
        except Exception:
            return sym, None

    with ThreadPoolExecutor(max_workers=len(syms)) as ex:
        pairs = list(ex.map(_fetch, syms))
    result = {sym: data for sym, data in pairs if data is not None}
    return result


def build_target_series(prices: pd.DataFrame, target_choice: str, custom: float, index: pd.Index) -> pd.Series:
    if target_choice == "Zero":
        return pd.Series(0.0, index=index)
    if target_choice == "Custom":
        return pd.Series(custom, index=index)
    if "^IRX" not in prices.columns or prices["^IRX"].dropna().empty:
        return pd.Series(0.0, index=index)
    irx = prices["^IRX"].reindex(index).ffill().bfill()
    return (irx / 100.0) / TRADING_DAYS


def _is_nan(x) -> bool:
    return x is None or (isinstance(x, float) and math.isnan(x))


def fmt_pct(x, digits: int = 2) -> str:
    return "n/a" if _is_nan(x) else f"{x*100:.{digits}f}%"


def fmt_num(x, digits: int = 3) -> str:
    return "n/a" if _is_nan(x) else f"{x:.{digits}f}"


def sig_badge(t_stat) -> str:
    if _is_nan(t_stat):
        return ""
    abs_t = abs(t_stat)
    if abs_t >= 2.0:
        return '<span class="badge badge-sig">sig ✓</span>'
    elif abs_t >= 1.5:
        return '<span class="badge badge-borderline">~sig</span>'
    else:
        return '<span class="badge badge-insig">n.s.</span>'


def card(label: str, value: str, sub: str, value_class: str = "", badge: str = "") -> str:
    return (
        f'<div class="metric-card">'
        f'<div class="metric-label">{label}{badge}</div>'
        f'<div class="metric-value {value_class}">{value}</div>'
        f'<div class="metric-sub">{sub}</div>'
        f"</div>"
    )


# ---------------------------------------------------------------------------
# Formatting helpers (also used by research context builder)
# ---------------------------------------------------------------------------
def fmt_large(x) -> str:
    if _is_nan(x) or x is None:
        return "n/a"
    if x >= 1e12:
        return f"${x/1e12:.2f}T"
    if x >= 1e9:
        return f"${x/1e9:.1f}B"
    if x >= 1e6:
        return f"${x/1e6:.0f}M"
    return f"${x:,.0f}"


def fmt_multiple(x, digits: int = 1) -> str:
    return "n/a" if (_is_nan(x) or x is None) else f"{x:.{digits}f}x"


def fmt_signed_pct(x, digits: int = 1) -> tuple[str, str]:
    if _is_nan(x) or x is None:
        return "n/a", ""
    color = "#16a34a" if x >= 0 else "#dc2626"
    return f"{x*100:+.{digits}f}%", color


def fmt_price(x) -> str:
    return "n/a" if (_is_nan(x) or x is None) else f"${x:,.2f}"


# ---------------------------------------------------------------------------
# (Chart helpers removed — rendering handled by embedded design HTML)
# ---------------------------------------------------------------------------
_CHART_LAYOUT_UNUSED = dict(
    plot_bgcolor="#faf7ef",
    paper_bgcolor="#f5f1e8",
    margin=dict(l=0, r=10, t=30, b=0),
    font=dict(family="IBM Plex Mono, monospace", size=11, color="#1a1f1a"),
)
_GRID = dict(gridcolor="#ddd6c2", gridwidth=1)
_ZERO_LINE = dict(line_dash="dash", line_color="#d8d1bd", line_width=1)

# Design palette
_C_POS     = "#2a6b54"   # forest green (pos)
_C_NEG     = "#c1471a"   # burnt orange (neg)
_C_PRIMARY = "#1f4e3d"   # forest green (primary)
_C_ACCENT  = "#c1471a"   # burnt orange (accent)
_C_MUTED   = "#918e7e"   # muted text


def plot_return_distribution(returns: pd.Series, target_avg: float) -> go.Figure:
    rmin, rmax = float(returns.min()), float(returns.max())
    bin_size = (rmax - rmin) / 50
    shared_bins = dict(start=rmin, end=rmax, size=bin_size)

    below = returns[returns < target_avg]
    above = returns[returns >= target_avg]

    fig = go.Figure()
    fig.add_trace(go.Histogram(
        x=below.values, name="Below target",
        marker_color=_C_NEG, opacity=0.82, xbins=shared_bins,
    ))
    fig.add_trace(go.Histogram(
        x=above.values, name="Above target",
        marker_color=_C_POS, opacity=0.82, xbins=shared_bins,
    ))

    mean_r = float(returns.mean())
    es5 = float(returns.quantile(0.05))
    fig.add_vline(x=mean_r, line_dash="dash", line_color=_C_MUTED, line_width=1.2,
                  annotation_text=f"μ {mean_r:.2%}", annotation_position="top right",
                  annotation_font_size=10)
    fig.add_vline(x=es5, line_dash="dot", line_color=_C_ACCENT, line_width=1.5,
                  annotation_text=f"ES5% {es5:.2%}", annotation_position="top left",
                  annotation_font_size=10)

    fig.update_layout(
        **_CHART_LAYOUT,
        barmode="overlay",
        height=300,
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        xaxis=dict(title="Daily Return", tickformat=".1%", **_GRID),
        yaxis=dict(title="Days", **_GRID),
    )
    return fig


def plot_drawdown(drawdown: pd.Series) -> go.Figure:
    min_idx = drawdown.idxmin()
    min_val = float(drawdown.min())
    min_idx_str = min_idx.strftime("%Y-%m-%d")

    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=drawdown.index.strftime("%Y-%m-%d"),
        y=drawdown.values * 100,
        fill="tozeroy",
        fillcolor="rgba(193, 71, 26, 0.10)",
        line=dict(color=_C_NEG, width=1.5),
        name="Drawdown",
        hovertemplate="%{x}: %{y:.2f}%<extra></extra>",
    ))
    fig.add_annotation(
        x=min_idx_str, y=min_val * 100,
        ax=0, ay=-50,
        text=f"MAX DD {min_val:.1%}",
        showarrow=True, arrowhead=2,
        arrowcolor=_C_NEG,
        bgcolor="#faf7ef", bordercolor=_C_NEG,
        borderwidth=1, borderpad=4,
        font=dict(size=10, family="IBM Plex Mono, monospace", color=_C_NEG),
    )
    fig.update_layout(
        **_CHART_LAYOUT,
        height=260,
        yaxis=dict(title="", ticksuffix="%",
                   tickfont=dict(family="IBM Plex Mono, monospace", size=10, color=_C_MUTED),
                   **_GRID),
        xaxis=dict(tickfont=dict(family="IBM Plex Mono, monospace", size=10, color=_C_MUTED), **_GRID),
        showlegend=False,
    )
    return fig


def plot_rolling_betas(
    rb: pd.DataFrame,
    bench_sym: str,
    sector_sym: str | None,
    has_sector: bool,
) -> tuple[go.Figure, go.Figure]:
    fig1 = go.Figure()
    fig1.add_trace(go.Scatter(
        x=rb.index, y=rb["bench"],
        name=f"β_mkt ({bench_sym})",
        line=dict(color=_C_PRIMARY, width=1.6),
        hovertemplate="%{x|%Y-%m-%d}: %{y:.3f}<extra>β_mkt</extra>",
    ))
    if has_sector and "sector" in rb.columns:
        fig1.add_trace(go.Scatter(
            x=rb.index, y=rb["sector"],
            name=f"β_sector ({sector_sym})",
            line=dict(color=_C_ACCENT, width=1.6),
            hovertemplate="%{x|%Y-%m-%d}: %{y:.3f}<extra>β_sector</extra>",
        ))
    fig1.add_hline(y=0, **_ZERO_LINE)
    fig1.add_hline(y=1, line_dash="dot", line_color="#d8d1bd", line_width=1,
                   annotation_text="β=1", annotation_font_size=9,
                   annotation_font_color=_C_MUTED)
    fig1.update_layout(
        **_CHART_LAYOUT,
        height=260,
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        yaxis=dict(title="Beta", **_GRID),
        xaxis=dict(**_GRID),
    )

    fig2 = go.Figure()
    fig2.add_trace(go.Scatter(
        x=rb.index, y=rb["dlvix"],
        name="β_vix",
        line=dict(color=_C_ACCENT, width=1.6),
        fill="tozeroy",
        fillcolor="rgba(193, 71, 26, 0.08)",
        hovertemplate="%{x|%Y-%m-%d}: %{y:.3f}<extra>β_vix</extra>",
    ))
    fig2.add_hline(y=0, **_ZERO_LINE)
    fig2.update_layout(
        **_CHART_LAYOUT,
        height=220,
        showlegend=False,
        yaxis=dict(title="β_vix", **_GRID),
        xaxis=dict(**_GRID),
    )
    return fig1, fig2


# ---------------------------------------------------------------------------
# Qualitative research — SEC EDGAR extraction (token-free), see quant_signals.
# Legacy AI pipeline removed: no API key, no per-run cost.
# ---------------------------------------------------------------------------
def export_framework_docx(ticker: str, long_name: str, sections: dict, timestamp: str) -> bytes:
    """Build and return a formatted .docx of the filing extracts."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()

    title_para = doc.add_heading(f"Qualitative Assessment: {long_name} ({ticker})", level=0)
    doc.add_paragraph(
        f"Date: {timestamp}  ·  Source: SEC EDGAR (10-K Items 1/1A/7, 8-K) — primary-source extracts, no AI"
    )
    doc.add_paragraph("")

    _SEC_EMOJIS = {"1": "🏢", "2": "🌐", "3": "🛡️", "4": "📈", "5": "👥", "6": "⚙️", "7": "🔢"}

    for section_name, section_text in sections.items():
        doc.add_heading(section_name, level=1)
        for block in section_text.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("- ") or block.startswith("* "):
                for line in block.splitlines():
                    clean = line.lstrip("-* ").strip()
                    if clean:
                        doc.add_paragraph(clean, style="List Bullet")
            else:
                doc.add_paragraph(block)
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# (interpret() removed — Interpretation component in design JS handles this)
# ---------------------------------------------------------------------------
def _interpret_unused(m: dict, ticker: str, bench_sym: str, sector_sym: str | None) -> str:
    bm_b = m["bm"]["beta"]
    bv_b = m["bv"]["beta"]
    if _is_nan(bm_b) or _is_nan(bv_b):
        return "Insufficient data for interpretation."
    parts = [
        f"On a +1% **{bench_sym}** day with other factors held fixed, "
        f"{ticker} moves about **{bm_b:.2f}%**."
    ]
    if m["has_sector"]:
        bs_b = m["bs"]["beta"]
        sec_dir = "rises" if bs_b > 0 else "falls"
        sig_s = (
            "statistically distinguishable from zero"
            if abs(m["bs"]["t"]) >= 1.96
            else "not statistically distinguishable from zero"
        )
        parts.append(
            f"For a 1% **{sector_sym}** sector excess return (i.e. sector "
            f"moves 1% beyond what the market predicts), {ticker} {sec_dir} by "
            f"**{abs(bs_b):.2f}%** — at t = {m['bs']['t']:.2f}, {sig_s}."
        )
    bv_dir = "falls" if bv_b < 0 else "rises"
    sig_v = (
        "statistically distinguishable from zero"
        if abs(m["bv"]["t"]) >= 1.96
        else "not statistically distinguishable from zero"
    )
    parts.append(
        f"On a 10% VIX rise (e.g. 15 → 16.5) with all else fixed, {ticker} "
        f"{bv_dir} by **{abs(bv_b)*0.10*100:.2f}%** — at t = {m['bv']['t']:.2f}, {sig_v}."
    )
    return " ".join(parts)


# ---------------------------------------------------------------------------
# Metrics
# ---------------------------------------------------------------------------
def compute_metrics(
    prices: pd.DataFrame,
    ticker: str,
    bench_sym: str,
    sector_sym: str | None,
    target_choice: str,
    custom_target: float,
):
    if bench_sym not in prices.columns or "^VIX" not in prices.columns:
        return None
    if ticker not in prices.columns or prices[ticker].dropna().empty:
        return None

    use_sector = sector_sym is not None and sector_sym in prices.columns
    cols_needed = [ticker, bench_sym] + ([sector_sym] if use_sector else [])

    px = prices.dropna(subset=[ticker])
    rets = px[cols_needed].pct_change().dropna(how="any")
    vix = px["^VIX"].reindex(rets.index).dropna()
    dlog_vix = np.log(vix).diff()

    parts = {"r": rets[ticker], "bench": rets[bench_sym], "dlvix": dlog_vix}
    if use_sector:
        parts["sector_raw"] = rets[sector_sym]
    df = pd.concat({k: v for k, v in parts.items()}, axis=1).dropna()
    if len(df) < 60:
        return None

    target = build_target_series(prices, target_choice, custom_target, df.index)

    # --- Semi-deviations ----------------------------------------------------
    excess = df["r"] - target
    n = len(excess)
    down_d = math.sqrt((np.minimum(excess, 0) ** 2).sum() / n)
    up_d = math.sqrt((np.maximum(excess, 0) ** 2).sum() / n)
    down_a = down_d * math.sqrt(TRADING_DAYS)
    up_a = up_d * math.sqrt(TRADING_DAYS)

    # --- Expected Shortfall (5%, daily) ------------------------------------
    var5 = df["r"].quantile(0.05)
    es5_daily = df["r"][df["r"] <= var5].mean()

    # --- Drawdown ----------------------------------------------------------
    cum = (1 + df["r"]).cumprod()
    dd = cum / cum.cummax() - 1.0
    max_dd = dd.min()

    # --- Sector residualization (Frisch-Waugh) -----------------------------
    factor_cols = ["bench"]
    if use_sector:
        X_sec = sm.add_constant(df[["bench"]])
        m_sec = sm.OLS(df["sector_raw"], X_sec).fit()
        df["sector"] = m_sec.resid
        factor_cols.append("sector")
    factor_cols.append("dlvix")

    # --- Joint regression with HAC SEs -------------------------------------
    X = sm.add_constant(df[factor_cols])
    y = df["r"]
    model = sm.OLS(y, X).fit(cov_type="HAC", cov_kwds={"maxlags": 5})

    def factor_block(name: str) -> dict:
        if name not in model.params.index:
            return {"beta": float("nan"), "t": float("nan"), "ci": [float("nan"), float("nan")]}
        return {
            "beta": float(model.params[name]),
            "t": float(model.tvalues[name]),
            "ci": [float(c) for c in model.conf_int().loc[name].tolist()],
        }

    bm = factor_block("bench")
    bs = factor_block("sector")
    bv = factor_block("dlvix")

    alpha_ann = float(model.params["const"]) * TRADING_DAYS
    sigma_eps_a = math.sqrt(model.scale) * math.sqrt(TRADING_DAYS)
    adj_r2 = float(model.rsquared_adj)

    # --- Other supporting numbers ------------------------------------------
    ann_return = (1 + df["r"].mean()) ** TRADING_DAYS - 1
    ann_vol = df["r"].std() * math.sqrt(TRADING_DAYS)
    rf_avg_daily = float(target.mean())
    ann_rf = (1 + rf_avg_daily) ** TRADING_DAYS - 1
    sortino = (ann_return - ann_rf) / down_a if down_a else float("nan")
    skew = df["r"].skew()
    kurt = df["r"].kurt()
    worst = df["r"].nsmallest(5)
    best = df["r"].nlargest(5)

    # --- Rolling betas -----------------------------------------------------
    try:
        roll = RollingOLS(y, X, window=60, min_nobs=60).fit()
        rolling_betas = roll.params.dropna(how="all")
    except Exception:
        rolling_betas = pd.DataFrame()

    down_m = down_d * math.sqrt(21)

    return {
        "n_obs": int(n),
        "first_date": df.index.min(),
        "last_date": df.index.max(),
        "downside_dev_d": down_d,
        "downside_dev_m": down_m,
        "downside_dev_a": down_a,
        "upside_dev_a": up_a,
        "es5_daily": es5_daily,
        "max_dd": max_dd,
        "alpha_ann": alpha_ann,
        "bm": bm, "bs": bs, "bv": bv,
        "sigma_eps_a": sigma_eps_a,
        "adj_r2": adj_r2,
        "ann_return": ann_return,
        "ann_vol": ann_vol,
        "ann_rf": ann_rf,
        "sortino": sortino,
        "skew": skew, "kurt": kurt,
        "worst": worst, "best": best,
        "returns": df["r"],
        "drawdown": dd,
        "rolling_betas": rolling_betas,
        "target_avg": rf_avg_daily,
        "has_sector": use_sector,
    }


# ---------------------------------------------------------------------------
# Buy signal backtest (VIX-filtered)
# ---------------------------------------------------------------------------
@st.cache_data(show_spinner=False, ttl=60 * 60 * 4)
def backtest_buy_signal(
    ticker: str,
    bench_sym: str,
    target_choice: str,
    custom_target: float,
    years: int,
    vix_threshold: float = 35.0,
) -> dict:
    prices = load_prices(ticker, years)
    if ticker not in prices.columns or "^VIX" not in prices.columns:
        return {}

    px  = prices[ticker].dropna()
    vix = prices["^VIX"].reindex(px.index).ffill()
    if len(px) < 504:
        return {}

    extreme = vix > vix_threshold

    rets   = px.pct_change()
    target = build_target_series(prices, target_choice, custom_target, rets.index)
    excess = rets - target

    excess_clean = excess.where(~extreme)
    px_clean     = px.where(~extreme)

    down_sq        = np.minimum(excess_clean, 0.0) ** 2
    rolling_dd_ann = (
        down_sq.rolling(252, min_periods=60).mean() ** 0.5
    ) * math.sqrt(TRADING_DAYS)

    rolling_high     = px_clean.rolling(252, min_periods=60).max()
    buy_price_series = rolling_high * (1 - rolling_dd_ann)

    signal_arr = (
        (px.values < buy_price_series.values)
        & (~extreme.values)
        & (~np.isnan(buy_price_series.values))
    )

    trigger_positions: list[int] = []
    last_pos = -999
    for i in range(len(signal_arr)):
        if signal_arr[i] and (i - last_pos) >= 252:
            trigger_positions.append(i)
            last_pos = i

    empty = {
        "nTriggers": 0, "triggers": [],
        "hitRate": None, "hitRateVsBench": None,
        "meanFwdReturn": None, "stdFwdReturn": None, "signalNoise": None,
        "vixThreshold": vix_threshold,
        "nExtremeDays": int(extreme.sum()), "nTotalDays": int(len(px)),
    }
    if not trigger_positions:
        return empty

    bench_s = prices[bench_sym].reindex(px.index).ffill() if bench_sym in prices.columns else None

    completed = []
    for pos in trigger_positions:
        fwd_pos = pos + 252
        if fwd_pos >= len(px):
            continue
        fwd_ret   = float(px.iloc[fwd_pos] / px.iloc[pos] - 1)
        bench_fwd = (
            float(bench_s.iloc[fwd_pos] / bench_s.iloc[pos] - 1)
            if bench_s is not None else None
        )
        completed.append({
            "date":        px.index[pos].strftime("%Y-%m-%d"),
            "entryPrice":  round(float(px.iloc[pos]), 2),
            "highPrice":   round(float(rolling_high.iloc[pos]), 2),
            "buyTarget":   round(float(buy_price_series.iloc[pos]), 2),
            "fwdReturn":   round(fwd_ret, 4),
            "benchReturn": round(bench_fwd, 4) if bench_fwd is not None else None,
            "beat":        (fwd_ret > bench_fwd) if bench_fwd is not None else None,
            "positive":    fwd_ret > 0,
        })

    if not completed:
        return empty

    fwd_rets  = [c["fwdReturn"] for c in completed]
    beat_list = [c["beat"] for c in completed if c["beat"] is not None]
    mean_fwd  = float(np.mean(fwd_rets))
    std_fwd   = float(np.std(fwd_rets, ddof=1)) if len(fwd_rets) > 1 else None
    hit_rate  = sum(r > 0 for r in fwd_rets) / len(fwd_rets)
    hit_bench = (sum(beat_list) / len(beat_list)) if beat_list else None
    sig_noise = (mean_fwd / std_fwd) if (std_fwd and std_fwd > 0) else None

    return {
        "nTriggers":      len(completed),
        "triggers":       completed,
        "hitRate":        round(hit_rate, 4),
        "hitRateVsBench": round(hit_bench, 4) if hit_bench is not None else None,
        "meanFwdReturn":  round(mean_fwd, 4),
        "stdFwdReturn":   round(std_fwd, 4) if std_fwd is not None else None,
        "signalNoise":    round(sig_noise, 3) if sig_noise is not None else None,
        "vixThreshold":   vix_threshold,
        "nExtremeDays":   int(extreme.sum()),
        "nTotalDays":     int(len(px)),
    }


# ---------------------------------------------------------------------------
# Cached compute_metrics — avoids re-running OLS on every Streamlit interaction
# ---------------------------------------------------------------------------
@st.cache_data(show_spinner=False, ttl=60 * 30)
def _compute_metrics_cached(
    ticker: str,
    bench_sym: str,
    sector_sym: str | None,
    target_choice: str,
    custom_target: float,
    years: int,
) -> dict | None:
    prices = load_prices(ticker, years)
    return compute_metrics(prices, ticker, bench_sym, sector_sym, target_choice, custom_target)


# ---------------------------------------------------------------------------
# Resolve benchmark and sector
# ---------------------------------------------------------------------------
if not ticker:
    st.info("Enter a ticker in the sidebar to begin.")
    st.stop()

with st.spinner(f"Loading {years}y of data for {ticker}…"):
    try:
        info = get_ticker_info(ticker)
        prices = load_prices(ticker, years)
    except Exception as exc:
        st.error(f"Failed to load data: {exc}")
        st.stop()

if ticker not in prices.columns or prices[ticker].dropna().empty:
    st.error(
        f"No price data returned for `{ticker}`. "
        "Check the symbol (e.g. `BRK-B` not `BRK.B`)."
    )
    st.stop()

market_cap = info.get("marketCap")
detected_tier = detect_cap_tier(market_cap)
detected_bench = CAP_TIERS[detected_tier][0]
detected_sector_name = info.get("sector")
detected_sector_etf = SECTOR_ETF.get(detected_sector_name) if detected_sector_name else None
long_name = info.get("longName") or ticker

# Resolve benchmark
if benchmark_choice == "Auto":
    bench_sym = detected_bench
    bench_tier_label = f"{detected_bench} ({detected_tier})"
    bench_source = "auto"
else:
    bench_sym = benchmark_choice.split("(")[1].rstrip(")")
    bench_tier_label = f"{bench_sym} (manual)"
    bench_source = "manual"

# Resolve sector
sector_sym: str | None = None
sector_source = "off"
if sector_mode == "Auto-detect":
    sector_sym = detected_sector_etf
    sector_source = "auto" if sector_sym else "auto (none detected)"
elif sector_mode == "Manual":
    sector_sym = manual_sector
    sector_source = "manual"

# Compute (cached — only re-runs when ticker/params change, not on every interaction)
m = _compute_metrics_cached(ticker, bench_sym, sector_sym, target_choice, custom_target, years)
if m is None:
    st.error("Not enough overlapping data to compute the model (need ≥ 60 obs).")
    st.stop()

bt = backtest_buy_signal(ticker, bench_sym, target_choice, custom_target, years)


# ---------------------------------------------------------------------------
# Research — token-free SEC EDGAR extraction (cached 12h per ticker)
# ---------------------------------------------------------------------------
try:
    _research = fetch_sec_filing_sections(ticker)
except Exception as _exc:
    _research = {"sections": [], "error": str(_exc)}
if _research.get("error"):
    st.sidebar.caption(f"SEC research: {_research['error']}")
_research_sections = _research.get("sections", [])
try:
    _research["drivers"] = extract_earnings_drivers(_research_sections)
except Exception:
    _research["drivers"] = {}

# Export filing extracts as .docx
if _research_sections:
    _parsed_for_export = {s["title"]: s["body"] for s in _research_sections}
    _docx = export_framework_docx(ticker, long_name, _parsed_for_export, date.today().isoformat())
    st.sidebar.download_button(
        label="⬇  Download Filing Extracts (.docx)",
        data=_docx,
        file_name=f"{ticker}_sec_extracts_{date.today().isoformat()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# ---------------------------------------------------------------------------
# Build + render the embedded design HTML
# ---------------------------------------------------------------------------
# Peer groups keyed by yfinance *industry* — true competitors that actually
# compete in the same product market. Preferred over the broad sector list so
# that, e.g., ASML is compared to other semiconductor-equipment makers (AMAT,
# LRCX, KLAC) rather than to Apple or Oracle.
INDUSTRY_PEERS = {
    # --- Technology ---
    "Semiconductors":                    ["NVDA", "AMD", "AVGO", "TSM", "INTC", "QCOM", "TXN", "MU", "ADI"],
    "Semiconductor Equipment & Materials": ["ASML", "AMAT", "LRCX", "KLAC", "TER", "ENTG", "ONTO", "ACLS"],
    "Software - Infrastructure":          ["MSFT", "ORCL", "PANW", "CRWD", "FTNT", "DDOG", "SNOW", "NET"],
    "Software - Application":             ["CRM", "ADBE", "INTU", "NOW", "SAP", "WDAY", "TEAM", "HUBS"],
    "Consumer Electronics":               ["AAPL", "SONY", "GPRO", "SONO"],
    "Information Technology Services":     ["ACN", "IBM", "FIS", "FISV", "INFY", "CTSH", "GIB", "EPAM"],
    "Communication Equipment":            ["CSCO", "ANET", "MSI", "JNPR", "NOK", "ERIC", "FFIV"],
    "Computer Hardware":                  ["DELL", "HPQ", "HPE", "NTAP", "WDC", "STX", "SMCI"],
    "Electronic Components":              ["APH", "TEL", "GLW", "JBL", "FLEX"],
    # --- Communication Services ---
    "Internet Content & Information":     ["GOOGL", "META", "PINS", "SNAP", "RDDT", "BIDU"],
    "Entertainment":                      ["NFLX", "DIS", "WBD", "PARA", "SPOT", "FOXA"],
    "Telecom Services":                   ["TMUS", "VZ", "T", "CMCSA", "CHTR"],
    # --- Consumer Cyclical ---
    "Internet Retail":                    ["AMZN", "BABA", "MELI", "EBAY", "ETSY", "CHWY"],
    "Auto Manufacturers":                 ["TSLA", "TM", "GM", "F", "STLA", "RIVN", "LCID"],
    "Restaurants":                        ["MCD", "SBUX", "CMG", "YUM", "DRI", "QSR"],
    "Footwear & Accessories":             ["NKE", "ADDYY", "DECK", "SKX", "CROX", "ONON"],
    "Home Improvement Retail":            ["HD", "LOW", "FND", "BLDR"],
    "Discount Stores":                    ["WMT", "COST", "TGT", "DG", "DLTR", "BJ"],
    # --- Healthcare ---
    "Drug Manufacturers - General":       ["LLY", "JNJ", "MRK", "ABBV", "PFE", "BMY", "NVS", "AZN"],
    "Biotechnology":                      ["AMGN", "GILD", "VRTX", "REGN", "BIIB", "MRNA"],
    "Healthcare Plans":                   ["UNH", "ELV", "CI", "HUM", "CNC", "CVS"],
    "Medical Devices":                    ["MDT", "ABT", "SYK", "BSX", "ISRG", "EW"],
    # --- Financials ---
    "Banks - Diversified":                ["JPM", "BAC", "WFC", "C", "USB", "PNC"],
    "Credit Services":                    ["V", "MA", "AXP", "PYPL", "COF", "DFS"],
    "Asset Management":                   ["BLK", "BX", "KKR", "APO", "BAM"],
    # --- Industrials ---
    "Aerospace & Defense":                ["BA", "RTX", "LMT", "GD", "NOC", "GE", "HWM"],
    "Farm & Heavy Construction Machinery": ["CAT", "DE", "PCAR", "CMI", "OSK"],
    "Integrated Freight & Logistics":     ["UPS", "FDX", "XPO", "CHRW"],
    # --- Energy ---
    "Oil & Gas Integrated":               ["XOM", "CVX", "SHEL", "BP", "TTE", "COP"],
    "Oil & Gas E&P":                      ["COP", "EOG", "OXY", "DVN", "FANG", "HES"],
    "Oil & Gas Equipment & Services":     ["SLB", "HAL", "BKR", "NOV", "FTI"],
}

# Default peer groups by yfinance *sector* — fallback when the industry isn't
# in INDUSTRY_PEERS and the user hasn't supplied a manual override.
SECTOR_PEERS = {
    "Technology":             ["AAPL", "MSFT", "NVDA", "AVGO", "ORCL", "CRM", "AMD"],
    "Communication Services": ["GOOGL", "META", "NFLX", "DIS", "TMUS", "T", "VZ"],
    "Consumer Cyclical":      ["AMZN", "TSLA", "HD", "MCD", "NKE", "SBUX", "LOW"],
    "Consumer Defensive":     ["WMT", "PG", "KO", "PEP", "COST", "MDLZ", "CL"],
    "Healthcare":             ["UNH", "JNJ", "LLY", "MRK", "ABBV", "PFE", "TMO"],
    "Financial Services":     ["JPM", "BAC", "WFC", "MA", "V", "GS", "MS"],
    "Industrials":            ["CAT", "HON", "GE", "BA", "UPS", "RTX", "DE"],
    "Energy":                 ["XOM", "CVX", "COP", "SLB", "EOG", "MPC", "PSX"],
    "Basic Materials":        ["LIN", "SHW", "FCX", "NEM", "APD", "ECL", "DOW"],
    "Real Estate":            ["PLD", "AMT", "EQIX", "SPG", "O", "PSA", "CCI"],
    "Utilities":              ["NEE", "DUK", "SO", "D", "AEP", "EXC", "SRE"],
}


@st.cache_data(show_spinner=False, ttl=60 * 60 * 4)
def get_peer_infos(peer_tickers: tuple[str, ...]) -> dict:
    """Fetch get_ticker_info() for each peer ticker, in parallel."""
    out: dict = {}
    with ThreadPoolExecutor(max_workers=6) as ex:
        futures = {ex.submit(get_ticker_info, t): t for t in peer_tickers}
        for fut in futures:
            t = futures[fut]
            try:
                out[t] = fut.result()
            except Exception:
                out[t] = {}
    return out


def historical_pe_average(prices, ticker: str, fund_history: dict | None,
                          years: int = 5):
    """The stock's own historical average P/E: for each of the last `years`
    fiscal years, divide that year's average closing price by that year's
    reported EPS, then average the yearly P/Es. Returns None if not derivable.

    Uses the daily price series + annual EPS from the fundamental-history feed,
    so it needs no extra data source.
    """
    try:
        if prices is None or ticker not in getattr(prices, "columns", []):
            return None
        px = prices[ticker].dropna()
        if px.empty:
            return None
        eps_pts = ((fund_history or {}).get("metrics", {})
                   .get("eps", {}) or {}).get("annual") or []
        yearly = []
        for p in eps_pts:
            d, v = p.get("d"), p.get("v")
            if not d or v in (None, 0):
                continue
            try:
                eps = float(v)
            except (TypeError, ValueError):
                continue
            if eps <= 0:
                continue
            yr = str(d)[:4]
            if not yr.isdigit():
                continue
            yr_px = px[px.index.year == int(yr)]
            if yr_px.empty:
                continue
            yearly.append(float(yr_px.mean()) / eps)
        yearly = yearly[:years]          # eps points are newest-first
        if yearly:
            return sum(yearly) / len(yearly)
    except Exception:
        pass
    return None


def resolve_peers(focus_ticker: str, sector: str | None, industry: str | None,
                  override_raw: str) -> tuple[list, str]:
    """Resolve the comparison set and a label describing its basis.

    Priority: manual override → industry competitors → sector fallback.
    Returns (peer_tickers, basis_label).
    """
    if override_raw and override_raw.strip():
        tickers = [t.strip().upper() for t in override_raw.split(",") if t.strip()]
        basis = "Custom list"
    elif industry and industry in INDUSTRY_PEERS:
        tickers = list(INDUSTRY_PEERS[industry])
        basis = f"Industry · {industry}"
    else:
        tickers = list(SECTOR_PEERS.get(sector or "", []))
        basis = f"Sector · {sector}" if tickers else "No peer group found"
    # Drop the focus ticker from the peer set, cap at 9 names.
    tickers = [t for t in tickers if t != focus_ticker.upper()][:9]
    return tickers, basis


_TAPE_SYMS = ("SPY", "QQQ", "IWM", "^VIX", "XLK", "XLF", "XLV", "XLE", "XLY", "IJH")
_WL_SYMS   = ("AAPL", "NVDA", "MSFT", "TSLA", "GOOGL", "META", "AMZN")
_WL_NAMES  = {
    "AAPL": "Apple Inc.", "NVDA": "NVIDIA Corp.", "MSFT": "Microsoft Corp.",
    "TSLA": "Tesla Inc.", "GOOGL": "Alphabet Inc.", "META": "Meta Platforms",
    "AMZN": "Amazon.com Inc.",
}

# Live quotes — 60-second TTL, covers market hours
_live_q    = get_live_quote(ticker)
_live_tape = get_live_tape_quotes(_TAPE_SYMS + _WL_SYMS)

# Watchlist with live prices
_WATCHLIST = []
for _sym in _WL_SYMS:
    _lq = _live_tape.get(_sym, {})
    _last  = _lq.get("last", 0)
    _prev  = _lq.get("prev", _last)
    _delta = (_last / _prev - 1) if _prev else 0.0
    _WATCHLIST.append({"sym": _sym, "name": _WL_NAMES[_sym],
                        "last": round(_last, 2), "delta": round(_delta, 6)})

with st.spinner(f"Rendering {ticker}…"):
    _js_data = to_js_data(m, prices, ticker, bench_sym, live_quote=_live_q, backtest=bt)
    _js_fund = to_js_fundamentals(info)
    _js_fund["ticker"] = ticker
    _js_tape = to_js_tape(prices, live_quotes=_live_tape)
    try:
        _js_quant = build_quant_payload(prices, ticker)
    except Exception as _exc:
        st.warning(f"Quant signals computation failed: {_exc}")
        _js_quant = {}
    try:
        _js_fund_hist = fetch_fundamental_history(ticker)
    except Exception as _exc:
        st.warning(f"Fundamental history fetch failed: {_exc}")
        _js_fund_hist = {}
    try:
        _js_credit = fetch_credit_metrics(ticker, info.get("marketCap"))
    except Exception:
        _js_credit = {}
    _js_credit["mentions"] = _research.get("ratingMentions", [])

    # Industry analysis — peer comparison payload
    try:
        _peer_syms, _peer_basis = resolve_peers(
            ticker, info.get("sector"), info.get("industry"), industry_peers_raw)
        _peer_infos = get_peer_infos(tuple(_peer_syms))
        _hist_pe = historical_pe_average(prices, ticker, _js_fund_hist)
        _js_industry = to_js_industry(ticker, info, _peer_infos,
                                      fund_history=_js_fund_hist,
                                      peer_basis=_peer_basis,
                                      extra_five_yr={"trailingPE": _hist_pe})
    except Exception as _exc:
        st.warning(f"Industry analysis computation failed: {_exc}")
        _js_industry = {}

    # Macro economics — G7 yields + GDP (cached inside macro_data)
    try:
        _js_macro = to_js_macro()
    except Exception as _exc:
        st.warning(f"Macro data computation failed: {_exc}")
        _js_macro = {}

    # Diagnostic line — shows whether each payload actually has data
    _fh_metrics = (_js_fund_hist or {}).get("metrics", {})
    _n_fund = sum(1 for v in _fh_metrics.values()
                  if v.get("annual") or v.get("quarterly"))
    _fh_src = (_js_fund_hist or {}).get("source", "none")
    _diag = (
        f"regime {'✓' if _js_quant.get('regime') else '✗'} · "
        f"momentum {'✓' if _js_quant.get('momentum') else '✗'} · "
        f"ewma {'✓' if _js_quant.get('ewma') else '✗'} · "
        f"iv/rv {'✓' if _js_quant.get('ivrv') else '✗'} · "
        f"fundamentals {_n_fund}/6 ({_fh_src}) · "
        f"credit {'✓' if _js_credit.get('zScore') is not None or _js_credit.get('totalDebt') is not None else '✗'} · "
        f"vix9d/3m {'✓' if ('^VIX9D' in prices.columns and '^VIX3M' in prices.columns) else '✗ (stale cache — press C → Clear cache)'}"
    )
    st.sidebar.caption(f"SIGNALS DIAG · {_diag}")
    _html = build_html(_js_data, _js_fund, _js_tape, _WATCHLIST, _research, ticker,
                       js_quant=_js_quant, js_fund_history=_js_fund_hist,
                       js_credit=_js_credit, js_industry=_js_industry,
                       js_macro=_js_macro)

def embed_html(html: str, height: int) -> None:
    """Embed a self-contained HTML document in an iframe.

    Prefers st.iframe (the supported API as of Streamlit 1.58); falls back to
    the deprecated st.components.v1.html on older Streamlit versions so the app
    still runs locally on an older install.
    """
    if hasattr(st, "iframe"):
        st.iframe(html, height=height)
    else:
        components.html(html, height=height, scrolling=False)


embed_html(_html, 2200)
